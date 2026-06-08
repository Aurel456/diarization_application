# src/exporter.py
import os
import re
import logging
import pandas as pd
from docx import Document
from docx.shared import Pt  # For font size if needed
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any, List, Optional
import time


def split_sentences_with_linebreaks(text: str, max_sentences: int = 6) -> str:
    if not text or not isinstance(text, str):
        return ""
    # Liste d'abréviations à ignorer
    abbreviations = ["M.", "Mme.", "Dr.", "Pr.", "etc.", "Mlle."]
    # On découpe en évitant les abréviations
    sentences = re.split(r"(?<!\bM)(?<!\bMme)(?<!\bDr)(?<!\bPr)(?<!\betc)\.\s+", text)
    grouped = []
    for i in range(0, len(sentences), max_sentences):
        chunk = ". ".join(sentences[i : i + max_sentences]).strip()
        if chunk and not chunk.endswith("."):
            chunk += "."
        grouped.append(chunk)
    return "\n\n".join(grouped)


def clean_before_export(text: Any) -> str:
    """
    Applies final cleaning rules (mostly regex substitutions) to text before export.

    Args:
        text: The text to clean (can be string, NaN, None).

    Returns:
        Cleaned text as a string.
    """
    if pd.isna(text) or not isinstance(text, str):
        return ""  # Return empty string for non-string inputs

    # Convert multiple spaces to single space
    res = " ".join(text.split())

    # Specific project/domain replacements (make configurable if needed)
    res = re.sub(r"Voici le texte corrigé\s?:\s?", "", res, flags=re.IGNORECASE)
    res = re.sub(
        r"Le texte corrigé est le suivant\s?:\s?", "", res, flags=re.IGNORECASE
    )
    res = res.replace('"', " ")  # Replace quotes with space
    # Add more replacements as needed, consider case sensitivity
    res = re.sub(r"\bXilo\b", "XYLO", res)  # Use word boundaries
    res = re.sub(r"\bxar\b", "CSAR", res)
    res = re.sub(r"\bXAR\b", "CSAR", res)
    res = re.sub(r"\des ténum\b", "DTNUM", res)
    res = re.sub(r"\bagraf\b", "AGRAF", res, flags=re.IGNORECASE)
    res = re.sub(r"\bDTNU(?!M)\b", "DTNUM", res)  # Negative lookahead

    # Artefacts structurels Cohere
    # Point entre deux minuscules : "de.souligner" → "de souligner"
    res = re.sub(r"(?<=[a-zéèêëàâùûüïîôçœæ])\.(?=[a-zéèêëàâùûüïîôçœæ])", " ", res)
    # Ponctuation parasite avant conjonction : "exister. et" / "organisé ? et" → "exister et"
    res = re.sub(r"\s*[.?!]\s+(et|ou|mais|donc|or|ni|car)\b", r" \1", res)
    # Préposition/article collé à un mot : "deCommunication" → "de communication"
    res = re.sub(
        r"\b(de|du|le|la|les|en|un|une|des|au|aux)([A-ZÉÈÊËÀÂÙÛÜÏÎÔÇ][a-zéèêëàâùûüïîôçœæ]+)\b",
        lambda m: m.group(1) + " " + m.group(2).lower(),
        res,
    )
    # Acronymes Cohere
    res = re.sub(r"\bXAM\b", "CSAM", res, flags=re.IGNORECASE)
    res = re.sub(r"\bADG\s+FIP\b", "DGFIP", res, flags=re.IGNORECASE)
    res = re.sub(r"d'ITRIX(?:\s+\d+)?", "dite Rixain", res, flags=re.IGNORECASE)

    # Trim final whitespace
    return res.strip()


def concatenate_texts(
    df: pd.DataFrame, speaker_col: str, text_col: str
) -> pd.DataFrame:
    """
    Concatenates text from consecutive rows if they have the same speaker.

    Args:
        df: DataFrame containing speaker and text columns.
        speaker_col: Name of the column identifying the speaker.
        text_col: Name of the column containing the text to concatenate.

    Returns:
        DataFrame with consecutive texts merged for the same speaker.
    """
    if df.empty:
        return df

    df = df.copy()
    # Create a grouping key that increments each time the speaker changes
    df["group"] = (df[speaker_col] != df[speaker_col].shift()).cumsum()

    # Group by this key and aggregate
    # Use 'first' for speaker and join texts with a space
    df_concat = (
        df.groupby("group")
        .agg(
            {
                speaker_col: "first",
                text_col: lambda x: " ".join(
                    x.astype(str).dropna()
                ),  # Join non-NA texts
            }
        )
        .reset_index(drop=True)
    )

    return df_concat


def save_dialogue_to_docx(
    df: pd.DataFrame,
    output_path: str,
    speaker_col: str = "global_speaker",
    text_col: str = "cleaned_transcription",
    dic_metadata: dict = None,
    sep: str = ": ",
    max_sentences_per_paragraph: int = 5,
    speaker_info: dict = None,
):
    document = Document()

    # --- Metadata (inchangé) ---
    if dic_metadata:
        p = document.add_paragraph(90 * "_")
        p.bold = True
        document.add_paragraph()
        for key, value in dic_metadata.items():
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run_key = p.add_run(f"{key}{sep}")
            run_key.bold = True
            p.add_run(str(value))
        document.add_paragraph()
        p = document.add_paragraph(90 * "_")
        p.bold = True
        document.add_paragraph()

    # --- Dialogue ---
    for _, row in df.iterrows():
        speaker_id = row.get(speaker_col)
        text = row.get(text_col)
        if pd.isna(speaker_id) or pd.isna(text):
            continue

        # Utiliser speaker_info pour formatter le nom si disponible.
        # Le LLM utilise "?" comme placeholder pour nom/prenom inconnus
        # (et "Inconnu" pour la fonction) — on filtre ces placeholders
        # pour garder l'ID du locuteur quand on n'a que la fonction.
        if speaker_info and speaker_id in speaker_info:
            info = speaker_info[speaker_id]
            prenom = (info.get("prenom") or "").strip()
            nom = (info.get("nom") or "").strip()
            fonction = (info.get("fonction") or "").strip()

            prenom_clean = "" if prenom == "?" else prenom
            nom_clean = "" if nom == "?" else nom
            fonction_clean = (
                "" if fonction.lower() in ("", "inconnu", "?") else fonction
            )
            name_part = f"{prenom_clean} {nom_clean}".strip()

            if name_part:
                speaker_display = name_part
                if fonction_clean:
                    speaker_display = f"{speaker_display} ({fonction_clean})"
            elif fonction_clean:
                # Pas de nom mais fonction connue → garder l'ID + (?) + fonction
                speaker_display = f"{speaker_id} (?) ({fonction_clean})"
            else:
                speaker_display = speaker_id
        else:
            speaker_display = speaker_id

        text = split_sentences_with_linebreaks(
            str(text), max_sentences=max_sentences_per_paragraph
        )

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_spk = p.add_run(f"{speaker_display}{sep}")
        run_spk.bold = True
        run_txt = p.add_run(text)

    try:
        document.save(output_path)
    except Exception as e:
        import logging

        logging.error(f"Failed to save DOCX: {e}", exc_info=True)


def export_results(
    data: pd.DataFrame,
    audio_paths: List[str],
    output_folder: str,
    speaker_info: dict = None,
) -> pd.DataFrame:
    """
    Prepares data and exports the final transcription dialogue to a DOCX file.

    Args:
        data: DataFrame with 'global_speaker' and 'cleaned_transcription' columns.
        audio_paths: List of paths to the original input audio files (used for naming).
        output_folder: Folder where the DOCX file will be saved.
        speaker_info: Optional dict mapping speaker_id to {nom, prenom, fonction}.

    Returns:
        The prepared DataFrame that was exported.
    """
    if (
        "global_speaker" not in data.columns
        or "cleaned_transcription" not in data.columns
    ):
        logging.error(
            "Export failed: DataFrame missing 'global_speaker' or 'cleaned_transcription'."
        )
        return pd.DataFrame()  # Return empty df

    if isinstance(audio_paths, list):
        # Get the base names of all audio paths and join them with a '+'
        base_audio_names = [
            os.path.splitext(os.path.basename(path))[0] for path in audio_paths
        ]
        base_audio_name = "+".join(base_audio_names)
    else:
        # Get the base name from the single audio path
        audio_name = os.path.basename(audio_paths)
        base_audio_name = os.path.splitext(audio_name)[0]

    logging.info("Preparing data for final export...")
    # Select relevant columns and handle potential NaNs in text
    export_df = data[["global_speaker", "cleaned_transcription"]].copy()
    export_df["cleaned_transcription"] = (
        export_df["cleaned_transcription"].fillna("").astype(str)
    )

    # Apply final cleaning rules
    export_df["cleaned_transcription"] = export_df["cleaned_transcription"].apply(
        clean_before_export
    )

    # Filter out rows where cleaning resulted in very short/empty text (e.g., just spaces or punctuation)
    # Adjust threshold '2' if needed
    export_df = export_df[export_df["cleaned_transcription"].str.strip().str.len() > 2]

    # Concatenate consecutive utterances from the same speaker
    logging.info("Concatenating consecutive utterances...")
    export_df_concatenated = concatenate_texts(
        export_df, "global_speaker", "cleaned_transcription"
    )

    # Prepare metadata
    num_repliques = len(export_df_concatenated)
    num_speakers = export_df_concatenated["global_speaker"].nunique()
    dic_metadata = {
        "Nom du fichier audio": base_audio_name,
        "Date de traitement": time.strftime("%Y-%m-%d %H:%M:%S"),
        # "Nom de la version": "Final Clustered & Cleaned", # Example version name
        "Nombre de répliques (après fusion)": f"{num_repliques}",
        "Nombre de speakers (estimé)": f"{num_speakers}",
    }
    logging.info(f"Export metadata: {dic_metadata}")

    # Define output path and save
    output_file = os.path.join(output_folder, f"Transcription_Final.docx")
    save_dialogue_to_docx(
        export_df_concatenated,
        output_file,
        speaker_col="global_speaker",
        text_col="cleaned_transcription",
        dic_metadata=dic_metadata,
        speaker_info=speaker_info,
    )

    return export_df_concatenated


def re_export_docx_with_labels(
    cleaned_data: pd.DataFrame,
    output_path: str,
    speaker_info: Optional[Dict[str, Dict[str, Any]]] = None,
    audio_label: str = "rapport",
) -> Optional[str]:
    """
    Regenerate a DOCX from an existing cleaned_data DataFrame with new speaker labels.

    Used when the user fills in names manually after the pipeline ran, or merges
    clusters. Avoids re-running transcription/clustering/cleaning.

    Args:
        cleaned_data: DataFrame containing 'global_speaker' and 'cleaned_transcription'.
        output_path:  Absolute path of the DOCX to write.
        speaker_info: Mapping speaker_id -> {"prenom", "nom", "fonction"}.
        audio_label:  Label written in the DOCX metadata.

    Returns:
        The output_path on success, None on failure.
    """
    if cleaned_data.empty or "global_speaker" not in cleaned_data.columns:
        logging.error("re_export_docx_with_labels: invalid DataFrame")
        return None

    df = cleaned_data[["global_speaker", "cleaned_transcription"]].copy()
    df["cleaned_transcription"] = df["cleaned_transcription"].fillna("").astype(str).apply(clean_before_export)
    df = df[df["cleaned_transcription"].str.strip().str.len() > 2]

    if df.empty:
        logging.warning("re_export_docx_with_labels: no content after cleaning")
        return None

    df_concat = concatenate_texts(df, "global_speaker", "cleaned_transcription")

    metadata = {
        "Nom du fichier audio": audio_label,
        "Date de traitement": time.strftime("%Y-%m-%d %H:%M:%S"),
        "Nombre de répliques (après fusion)": f"{len(df_concat)}",
        "Nombre de speakers (estimé)": f"{df_concat['global_speaker'].nunique()}",
        "Origine": "Regénéré avec labels manuels / fusion de clusters",
    }

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    save_dialogue_to_docx(
        df_concat,
        output_path,
        speaker_col="global_speaker",
        text_col="cleaned_transcription",
        dic_metadata=metadata,
        speaker_info=speaker_info,
    )
    return output_path if os.path.exists(output_path) else None


def apply_speaker_mapping(
    cleaned_data: pd.DataFrame,
    speaker_mapping: Dict[str, str],
) -> pd.DataFrame:
    """
    Apply a speaker_id -> new_speaker_id mapping to the 'global_speaker' column.

    Used by the cluster merge/split UI: e.g. {"Speaker_02": "Speaker_00"} merges
    cluster 02 into cluster 00. Returns a new DataFrame.
    """
    if "global_speaker" not in cleaned_data.columns:
        return cleaned_data
    out = cleaned_data.copy()
    out["global_speaker"] = out["global_speaker"].map(lambda s: speaker_mapping.get(s, s))
    return out
